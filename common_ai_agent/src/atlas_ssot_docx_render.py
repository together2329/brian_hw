"""SSOT DOCX render — split-out chunk (Phase 26 of refactor/atlas-modular).

Top 8 biggest functions from atlas_ssot_docx.py moved here to keep both
files under 1000 lines. Lazy hydration of helpers from sibling.
"""
from __future__ import annotations

from __future__ import annotations
import html as _html
import io
import json as _json
import os
import re
import re as _re
from pathlib import Path
from typing import Any, Optional
import yaml as _yaml
from src.atlas_ssot_export import (
    _SSOT_EXPORT_SECTION_ORDER,
    _ssot_md_scalar,
    _ssot_md_is_short_scalar,
    _ssot_section_is_empty,
    _ssot_field_bit_info,
)


def _import_docx_helpers():
    g = globals()
    if g.get("_HELPERS_HYDRATED"): return
    from src import atlas_ssot_docx as _d
    for name in (
        "_SSOT_DOCX_APPENDIX_KEYS",
        "_ssot_docx_add_list_of_figures",
        "_ssot_docx_add_list_of_tables",
        "_ssot_docx_add_toc",
        "_ssot_docx_apply_heading_numbering",
        "_ssot_docx_caption",
        "_ssot_docx_cover_page",
        "_ssot_docx_dict_block",
        "_ssot_docx_list_of_dicts",
        "_ssot_docx_page_break",
        "_ssot_docx_port_rows",
        "_ssot_docx_render_error_handling",
        "_ssot_docx_render_features",
        "_ssot_docx_render_function_description",
        "_ssot_docx_render_interrupts",
        "_ssot_docx_render_parameters",
        "_ssot_docx_render_port_buckets",
        "_ssot_docx_render_programming_model",
        "_ssot_docx_render_revision_history",
        "_ssot_docx_set_footer",
        "_ssot_docx_table_from_rows",
        "_ssot_docx_yaml_block",
    ):
        if hasattr(_d, name): g[name] = getattr(_d, name)
    g["_HELPERS_HYDRATED"] = True

_import_docx_helpers()

def _ssot_docx_render_section(doc: Any, key: str, value: Any) -> None:
    if key == "io_list" and isinstance(value, dict):
        cd = value.get("clock_domains")
        if cd:
            doc.add_heading("Clock Domains", level=2)
            if isinstance(cd, list) and cd and all(isinstance(r, dict) for r in cd):
                _ssot_docx_list_of_dicts(doc, cd)
            else:
                _ssot_docx_yaml_block(doc, cd)
        rst = value.get("resets")
        if rst:
            doc.add_heading("Resets", level=2)
            if isinstance(rst, list) and rst and all(isinstance(r, dict) for r in rst):
                _ssot_docx_list_of_dicts(doc, rst)
            else:
                _ssot_docx_yaml_block(doc, rst)
        ifs = value.get("interfaces")
        if isinstance(ifs, list):
            doc.add_heading("Interfaces", level=2)
            for iface in ifs:
                if not isinstance(iface, dict):
                    _ssot_docx_yaml_block(doc, iface)
                    continue
                doc.add_heading(str(iface.get("name") or "(unnamed)"), level=3)
                meta = {k: v for k, v in iface.items() if k != "ports"}
                _ssot_docx_dict_block(doc, meta)
                ports = iface.get("ports")
                if isinstance(ports, list) and ports:
                    _ssot_docx_list_of_dicts(doc, ports)
        leftover = {k: v for k, v in value.items()
                    if k not in ("clock_domains", "resets", "interfaces")}
        if leftover:
            _ssot_docx_dict_block(doc, leftover)
        return

    if key == "registers" and isinstance(value, dict):
        cfg = value.get("config")
        if isinstance(cfg, dict) and cfg:
            doc.add_heading("Config", level=2)
            _ssot_docx_dict_block(doc, cfg)
        reg_list = value.get("register_list")
        if isinstance(reg_list, list) and reg_list:
            doc.add_heading("Register List", level=2)
            _ssot_docx_list_of_dicts(doc, reg_list)
            for reg in reg_list:
                if not isinstance(reg, dict):
                    continue
                fields = reg.get("fields")
                if isinstance(fields, list) and fields:
                    doc.add_heading(
                        f"{reg.get('name') or '(unnamed reg)'} — fields",
                        level=3,
                    )
                    _ssot_docx_list_of_dicts(doc, fields)
        leftover = {k: v for k, v in value.items() if k not in ("config", "register_list")}
        if leftover:
            _ssot_docx_dict_block(doc, leftover)
        return

    if key == "features" and isinstance(value, list):
        for feat in value:
            if not isinstance(feat, dict):
                _ssot_docx_yaml_block(doc, feat)
                continue
            doc.add_heading(str(feat.get("name") or "(unnamed feature)"), level=2)
            rest = {k: v for k, v in feat.items() if k != "name"}
            _ssot_docx_dict_block(doc, rest)
        return

    if isinstance(value, dict):
        _ssot_docx_dict_block(doc, value)
        return
    if isinstance(value, list):
        if value and all(isinstance(r, dict) for r in value):
            _ssot_docx_list_of_dicts(doc, value)
            return
        if value and all(_ssot_md_is_short_scalar(r) for r in value):
            for item in value:
                doc.add_paragraph(_ssot_md_scalar(item), style="List Bullet")
            return
        _ssot_docx_yaml_block(doc, value)
        return
    if _ssot_md_is_short_scalar(value):
        doc.add_paragraph(_ssot_md_scalar(value))
        return
    _ssot_docx_yaml_block(doc, value)

def _ssot_docx_render_io_list(doc: Any, value: Any) -> None:
    """Clock/Reset/Interface (with direction-grouped port tables) signal description."""
    if isinstance(value, list):
        bucketed = _ssot_docx_port_rows(value)
        if bucketed:
            _ssot_docx_render_port_buckets(doc, bucketed)
            return
        _ssot_docx_yaml_block(doc, value)
        return

    if not isinstance(value, dict):
        _ssot_docx_yaml_block(doc, value)
        return

    cd = value.get("clock_domains") or value.get("clocks")
    if isinstance(cd, list) and cd:
        doc.add_heading("Clock domains", level=2)
        rows = []
        for entry in cd:
            if not isinstance(entry, dict):
                continue
            rows.append([
                _ssot_md_scalar(entry.get("name") or ""),
                _ssot_md_scalar(entry.get("frequency_mhz") or entry.get("frequency") or ""),
                _ssot_md_scalar(entry.get("source") or ""),
                _ssot_md_scalar(entry.get("description") or ""),
            ])
        if rows:
            _ssot_docx_caption(doc, "Table", "Clock domains")
            _ssot_docx_table_from_rows(doc, ["Domain", "Frequency (MHz)", "Source", "Description"], rows)

    rst = value.get("resets")
    if isinstance(rst, list) and rst:
        doc.add_heading("Resets", level=2)
        rows = []
        for entry in rst:
            if not isinstance(entry, dict):
                continue
            rows.append([
                _ssot_md_scalar(entry.get("name") or ""),
                _ssot_md_scalar(entry.get("polarity") or ""),
                _ssot_md_scalar(entry.get("sync_async") or entry.get("type") or ""),
                _ssot_md_scalar(entry.get("source") or ""),
                _ssot_md_scalar(entry.get("description") or ""),
            ])
        if rows:
            _ssot_docx_caption(doc, "Table", "Reset sources")
            _ssot_docx_table_from_rows(doc, ["Reset", "Polarity", "Sync / Async", "Source", "Description"], rows)

    ifs = value.get("interfaces")
    if isinstance(ifs, list) and ifs:
        doc.add_heading("Interfaces", level=2)
        for iface in ifs:
            if not isinstance(iface, dict):
                _ssot_docx_yaml_block(doc, iface)
                continue
            iname = str(iface.get("name") or "(unnamed)")
            doc.add_heading(iname, level=3)
            description = str(iface.get("description") or "").strip()
            if description:
                doc.add_paragraph(description)
            iface_meta_keys = [
                ("type", "Type"),
                ("role", "Role"),
                ("protocol", "Protocol"),
                ("data_width", "Data width"),
                ("addr_width", "Address width"),
                ("clock", "Clock domain"),
                ("reset", "Reset domain"),
            ]
            meta_rows = []
            for k, label in iface_meta_keys:
                v = iface.get(k)
                if v not in (None, "", [], {}):
                    meta_rows.append([label, _ssot_md_scalar(v)])
            if meta_rows:
                _ssot_docx_table_from_rows(doc, ["Property", "Value"], meta_rows)
            ports = iface.get("ports")
            if isinstance(ports, list) and ports:
                bucketed = _ssot_docx_port_rows(ports)
                _ssot_docx_caption(doc, "Table", f"Signal description — {iname}")
                _ssot_docx_render_port_buckets(doc, bucketed)

    flat_ports = value.get("ports")
    if isinstance(flat_ports, list) and flat_ports and not ifs:
        doc.add_heading("Signals", level=2)
        bucketed = _ssot_docx_port_rows(flat_ports)
        _ssot_docx_render_port_buckets(doc, bucketed)

    leftover = {
        k: v for k, v in value.items()
        if k not in ("clock_domains", "clocks", "resets", "interfaces", "ports")
        and not _ssot_section_is_empty(v)
    }
    if leftover:
        doc.add_heading("Additional I/O metadata", level=2)
        _ssot_docx_dict_block(doc, leftover)

def _ssot_docx_render_fsm(doc: Any, value: Any) -> None:
    """State list + transition table per machine."""
    machines: list[dict] = []
    if isinstance(value, dict):
        ml = value.get("machines") or value.get("fsm_list") or value.get("list")
        if isinstance(ml, list):
            machines = [m for m in ml if isinstance(m, dict)]
        else:
            # Pattern: {<machine_name>: {states: [...], transitions: [...]}}
            for k, v in value.items():
                if isinstance(v, dict) and (v.get("states") or v.get("transitions")):
                    machines.append({"name": k, **v})
    elif isinstance(value, list):
        machines = [m for m in value if isinstance(m, dict)]
    if not machines:
        _ssot_docx_yaml_block(doc, value)
        return
    for m in machines:
        name = str(m.get("name") or "(unnamed FSM)")
        doc.add_heading(name, level=2)
        description = str(m.get("description") or m.get("purpose") or "").strip()
        if description:
            doc.add_paragraph(description)
        meta_keys = [
            ("reset_state", "Reset state"),
            ("encoding", "Encoding"),
            ("clock", "Clock domain"),
            ("reset", "Reset domain"),
            ("illegal_recovery", "Illegal recovery"),
        ]
        meta_rows = []
        for k, label in meta_keys:
            v = m.get(k)
            if v not in (None, "", [], {}):
                meta_rows.append([label, _ssot_md_scalar(v)])
        if meta_rows:
            _ssot_docx_table_from_rows(doc, ["Property", "Value"], meta_rows)
        states = m.get("states")
        if isinstance(states, list) and states:
            doc.add_heading("States", level=3)
            rows = []
            for s in states:
                if isinstance(s, dict):
                    rows.append([
                        _ssot_md_scalar(s.get("id") or s.get("name") or ""),
                        _ssot_md_scalar(s.get("label") or s.get("name") or ""),
                        _ssot_md_scalar(s.get("encoding") or ""),
                        "yes" if s.get("reset") else "",
                        _ssot_md_scalar(s.get("description") or s.get("action") or ""),
                    ])
                else:
                    rows.append(["", _ssot_md_scalar(s), "", "", ""])
            _ssot_docx_table_from_rows(doc, ["ID", "State", "Encoding", "Reset", "Description / Action"], rows)
        transitions = m.get("transitions")
        if isinstance(transitions, list) and transitions:
            doc.add_heading("Transitions", level=3)
            rows = []
            for tr in transitions:
                if isinstance(tr, dict):
                    rows.append([
                        _ssot_md_scalar(tr.get("from") or tr.get("source") or tr.get("current") or ""),
                        _ssot_md_scalar(tr.get("event") or tr.get("condition") or tr.get("trigger") or ""),
                        _ssot_md_scalar(tr.get("to") or tr.get("dest") or tr.get("target") or tr.get("next") or ""),
                        _ssot_md_scalar(tr.get("action") or tr.get("output") or ""),
                        _ssot_md_scalar(tr.get("description") or ""),
                    ])
                else:
                    rows.append(["", _ssot_md_scalar(tr), "", "", ""])
            _ssot_docx_table_from_rows(
                doc, ["Current", "Event / Condition", "Next", "Action / Output", "Description"], rows,
            )
        outputs = m.get("outputs") or m.get("output_rules")
        if outputs:
            doc.add_heading("Output rules", level=3)
            if isinstance(outputs, list) and outputs and all(isinstance(o, dict) for o in outputs):
                _ssot_docx_list_of_dicts(doc, outputs)
            else:
                _ssot_docx_yaml_block(doc, outputs)

def _ssot_docx_render_cycle_model(doc: Any, value: Any) -> None:
    """Cycle contract prose + Latency table + Handshake table + Pipeline table."""
    if not isinstance(value, dict):
        _ssot_docx_yaml_block(doc, value)
        return
    purpose = str(value.get("purpose") or value.get("description") or "").strip()
    if purpose:
        doc.add_paragraph(purpose)
    contract_keys = [
        ("clock", "Clock domain"),
        ("reset", "Reset domain"),
        ("assertion", "Reset assertion"),
        ("deassertion", "Reset deassertion"),
        ("scope", "Scope"),
    ]
    contract_rows = []
    for k, label in contract_keys:
        v = value.get(k)
        if v not in (None, "", [], {}):
            contract_rows.append([label, _ssot_md_scalar(v)])
    if contract_rows:
        doc.add_heading("Cycle contract", level=2)
        _ssot_docx_table_from_rows(doc, ["Property", "Value"], contract_rows)
    latency = value.get("latency") or value.get("latencies")
    if isinstance(latency, dict) and latency:
        doc.add_heading("Latency", level=2)
        rows = []
        for path, info in latency.items():
            if isinstance(info, dict):
                rows.append([
                    str(path),
                    _ssot_md_scalar(info.get("min_cycles") or info.get("min") or ""),
                    _ssot_md_scalar(info.get("max_cycles") or info.get("max") or ""),
                    _ssot_md_scalar(info.get("description") or info.get("notes") or ""),
                ])
            else:
                rows.append([str(path), "", "", _ssot_md_scalar(info)])
        _ssot_docx_table_from_rows(doc, ["Path", "Min cycles", "Max cycles", "Description"], rows)
    elif isinstance(latency, list) and latency:
        doc.add_heading("Latency", level=2)
        rows = []
        for ent in latency:
            if isinstance(ent, dict):
                rows.append([
                    _ssot_md_scalar(ent.get("path") or ent.get("name") or ent.get("from_to") or ""),
                    _ssot_md_scalar(ent.get("min_cycles") or ent.get("min") or ""),
                    _ssot_md_scalar(ent.get("max_cycles") or ent.get("max") or ""),
                    _ssot_md_scalar(ent.get("description") or ent.get("notes") or ""),
                ])
        _ssot_docx_table_from_rows(doc, ["Path", "Min cycles", "Max cycles", "Description"], rows)
    handshake = value.get("handshake_rules") or value.get("handshake")
    if isinstance(handshake, list) and handshake:
        doc.add_heading("Handshake rules", level=2)
        rows = []
        for ent in handshake:
            if isinstance(ent, dict):
                rows.append([
                    _ssot_md_scalar(ent.get("signal") or ent.get("name") or ""),
                    _ssot_md_scalar(ent.get("rule") or ent.get("contract") or ent.get("description") or ""),
                    _ssot_md_scalar(ent.get("notes") or ""),
                ])
        _ssot_docx_table_from_rows(doc, ["Signal", "Rule", "Notes"], rows)
    pipeline = value.get("pipeline") or value.get("stages")
    if isinstance(pipeline, list) and pipeline:
        doc.add_heading("Pipeline", level=2)
        rows = []
        for ent in pipeline:
            if isinstance(ent, dict):
                rows.append([
                    _ssot_md_scalar(ent.get("stage") or ent.get("name") or ent.get("id") or ""),
                    _ssot_md_scalar(ent.get("cycle") or ent.get("phase") or ""),
                    _ssot_md_scalar(ent.get("action") or ent.get("description") or ""),
                    _ssot_md_scalar(ent.get("ready") or ""),
                    _ssot_md_scalar(ent.get("notes") or ""),
                ])
        _ssot_docx_table_from_rows(doc, ["Stage", "Cycle", "Action", "Ready / Backpressure", "Notes"], rows)
    leftover = {
        k: v for k, v in value.items()
        if k not in dict(contract_keys)
        and k not in ("latency", "latencies", "handshake_rules", "handshake", "pipeline", "stages", "purpose", "description")
        and not _ssot_section_is_empty(v)
    }
    if leftover:
        doc.add_heading("Additional cycle-model fields", level=2)
        _ssot_docx_dict_block(doc, leftover)

def _ssot_docx_block_diagram_png(data: dict):
    """Render an Andes-style block diagram PNG (BytesIO) via matplotlib.

    Returns None when matplotlib is not importable; the caller then
    emits only the text-and-table representation.
    """
    try:
        import matplotlib  # type: ignore
        matplotlib.use("Agg", force=True)
        import matplotlib.pyplot as plt  # type: ignore
        import matplotlib.patches as mpatches  # type: ignore
    except Exception:
        return None
    import io as _io

    top = data.get("top_module") if isinstance(data.get("top_module"), dict) else {}
    submods_raw = data.get("sub_modules") if isinstance(data.get("sub_modules"), list) else []
    submods = [sm for sm in submods_raw if isinstance(sm, (dict, str))]
    io_list = data.get("io_list") if isinstance(data.get("io_list"), dict) else {}
    interfaces = io_list.get("interfaces") if isinstance(io_list, dict) else None
    if not isinstance(interfaces, list):
        interfaces = []

    fig_w = 9.5
    fig_h = 6.0 if submods else 4.0
    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, fig_h)
    ax.axis("off")

    outer_x, outer_w = 1.5, 7.0
    outer_y, outer_h = 0.6, fig_h - 1.0
    outer = mpatches.FancyBboxPatch(
        (outer_x, outer_y), outer_w, outer_h,
        boxstyle="round,pad=0.02",
        edgecolor="#4a6cf7", facecolor="#f4f7ff", lw=2,
    )
    ax.add_patch(outer)
    top_name = str(top.get("name") or "TOP")
    top_role = str(top.get("type") or top.get("role") or "").strip()
    ax.text(
        outer_x + outer_w / 2, outer_y + outer_h - 0.35,
        top_name + (f"   ({top_role})" if top_role else ""),
        ha="center", va="center", fontsize=14, fontweight="bold",
    )

    if submods:
        cols = min(3, len(submods))
        rows = (len(submods) + cols - 1) // cols
        inner_w = outer_w - 0.6
        inner_h = outer_h - 1.4
        cell_w = inner_w / cols
        cell_h = inner_h / rows
        for i, sm in enumerate(submods):
            r = i // cols
            c = i % cols
            x = outer_x + 0.3 + c * cell_w
            y = outer_y + 0.4 + (rows - 1 - r) * cell_h
            sub = mpatches.FancyBboxPatch(
                (x + 0.08, y + 0.08), cell_w - 0.16, cell_h - 0.16,
                boxstyle="round,pad=0.02",
                edgecolor="#4a6cf7", facecolor="#ffffff", lw=1.2,
            )
            ax.add_patch(sub)
            name = (sm.get("name") if isinstance(sm, dict) else str(sm)) or "?"
            desc = ""
            if isinstance(sm, dict):
                desc = str(sm.get("description") or sm.get("implementation") or "")[:64]
            ax.text(
                x + cell_w / 2, y + cell_h / 2 + (0.15 if desc else 0),
                name, ha="center", va="center",
                fontsize=10, fontweight="bold", family="monospace",
            )
            if desc:
                ax.text(
                    x + cell_w / 2, y + cell_h / 2 - 0.2,
                    desc, ha="center", va="center",
                    fontsize=7, color="#5a6a8f",
                )

    valid_ifs = [i for i in interfaces if isinstance(i, dict)][:4]
    in_color, out_color = "#4a6cf7", "#7a4af7"
    for idx, ifc in enumerate(valid_ifs):
        name = str(ifc.get("name") or "?")
        ifc_type = str(ifc.get("type") or "").strip()
        role = str(ifc.get("role") or "").lower()
        inward = role in {"slave", "sink", "in", "input", ""}
        side_left = idx % 2 == 0
        y_pos = outer_y + outer_h - 0.9 - (idx // 2) * 1.6
        if y_pos < outer_y + 0.4:
            y_pos = outer_y + outer_h / 2
        if side_left:
            arrow = mpatches.FancyArrowPatch(
                (0.4, y_pos) if inward else (outer_x + 0.1, y_pos),
                (outer_x + 0.1, y_pos) if inward else (0.4, y_pos),
                arrowstyle="->", mutation_scale=18,
                color=in_color if inward else out_color, lw=1.6,
            )
            ax.text(0.55, y_pos + 0.22, name + (f"\n({ifc_type})" if ifc_type else ""),
                    ha="center", va="bottom", fontsize=9, color="#1f2a44")
        else:
            arrow = mpatches.FancyArrowPatch(
                (9.6, y_pos) if inward else (outer_x + outer_w - 0.1, y_pos),
                (outer_x + outer_w - 0.1, y_pos) if inward else (9.6, y_pos),
                arrowstyle="->", mutation_scale=18,
                color=in_color if inward else out_color, lw=1.6,
            )
            ax.text(9.45, y_pos + 0.22, name + (f"\n({ifc_type})" if ifc_type else ""),
                    ha="center", va="bottom", fontsize=9, color="#1f2a44")
        ax.add_patch(arrow)

    buf = _io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=160, facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf

def _ssot_docx_render_block_diagram(doc: Any, data: dict) -> None:
    """Block-diagram chapter body — real figure first, submodule table beneath as key."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Inches, Pt  # type: ignore

    fig_buf = _ssot_docx_block_diagram_png(data)
    if fig_buf is not None:
        doc.add_picture(fig_buf, width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _ssot_docx_caption(doc, "Figure", "Block diagram")

    top = data.get("top_module") if isinstance(data.get("top_module"), dict) else {}
    top_name = str(top.get("name") or "TOP").strip() or "TOP"
    top_role = str(top.get("type") or top.get("role") or "").strip()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(top_name + (f"  ({top_role})" if top_role else ""))
    title_run.bold = True
    title_run.font.size = Pt(13)

    submods = data.get("sub_modules")
    if isinstance(submods, list) and submods:
        cols = min(3, len(submods))
        rows = (len(submods) + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        try:
            table.style = "Light Grid Accent 1"
        except Exception:
            pass
        for idx, sm in enumerate(submods):
            r = idx // cols
            c = idx % cols
            cell = table.rows[r].cells[c]
            if isinstance(sm, dict):
                name = str(sm.get("name") or sm.get("module") or "?").strip()
                desc = str(sm.get("description") or sm.get("implementation") or "").strip()
            else:
                name = str(sm)
                desc = ""
            cell.text = name
            for run in cell.paragraphs[0].runs:
                run.bold = True
            if desc:
                p = cell.add_paragraph()
                p.add_run(desc[:160]).italic = True
        # Center each cell
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        caption = doc.add_paragraph()
        cap = caption.add_run("(no sub-modules declared — single-block design)")
        cap.italic = True
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # External interface chips
    io_list = data.get("io_list") if isinstance(data.get("io_list"), dict) else {}
    interfaces = io_list.get("interfaces") if isinstance(io_list, dict) else None
    if isinstance(interfaces, list) and interfaces:
        legend = doc.add_paragraph()
        legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
        legend.add_run("External interfaces: ").italic = True
        legend.add_run(", ".join(
            str(i.get("name") or "?") + (f" ({i.get('type')})" if isinstance(i, dict) and i.get("type") else "")
            for i in interfaces if isinstance(i, dict)
        ))

def _ssot_docx_render_programming_sequence(doc: Any, data: dict) -> None:
    """Andes §5 layout — Setup / TX / RX style sub-sections derived from SSOT.

    Heuristics:
      * "IP Setup": find Enable / Configure registers, emit a numbered
        step list with the relevant register writes.
      * Scenarios: if cycle_model.scenarios / function_model.scenarios
        exist, one H3 per scenario walking cycle-by-cycle through the
        steps with action prose.
    """
    cycle_model = data.get("cycle_model") if isinstance(data.get("cycle_model"), dict) else {}
    fn_model = data.get("function_model") if isinstance(data.get("function_model"), dict) else {}
    scenarios = []
    for src in (cycle_model.get("scenarios"), fn_model.get("scenarios")):
        if isinstance(src, list):
            scenarios.extend(s for s in src if isinstance(s, dict))

    # 5.1 IP Setup
    registers = data.get("registers")
    reg_list = []
    if isinstance(registers, dict) and isinstance(registers.get("register_list"), list):
        reg_list = [r for r in registers["register_list"] if isinstance(r, dict)]
    elif isinstance(registers, list):
        reg_list = [r for r in registers if isinstance(r, dict)]
    enable_targets = []
    configure_targets = []
    for reg in reg_list:
        rn = str(reg.get("name") or "")
        fields = reg.get("fields") if isinstance(reg.get("fields"), list) else []
        en_field = next((f for f in fields if isinstance(f, dict)
                         and re.search(r"\ben\b|enable", str(f.get("name") or ""), re.I)), None)
        if en_field:
            enable_targets.append((rn, en_field.get("name"), en_field.get("description")))
        elif re.match(r"^(cr|ctrl|control|cfg|config|mode)\b", rn, re.I):
            enable_targets.append((rn, None, reg.get("description")))
        if re.match(r"^(dlr|brr|baud|div|mr|prescaler|cfg)", rn, re.I) and not any(rn == e[0] for e in enable_targets):
            configure_targets.append((rn, reg.get("offset"), reg.get("description")))
    if enable_targets or configure_targets:
        doc.add_heading("IP Setup", level=2)
        steps = []
        for rn, fld, desc in enable_targets[:2]:
            steps.append(f"Write {rn}.{fld}=1 to enable the module." if fld else f"Configure {rn} (the primary control register) to bring the module out of reset.")
        for rn, off, desc in configure_targets[:3]:
            label = f"{rn}" + (f" (@ {off})" if off else "")
            steps.append(f"Program {label} with the operating parameters required for your use case.")
        if not steps:
            steps.append("Refer to the Register Description chapter for the bring-up sequence specific to this IP.")
        for i, s in enumerate(steps, start=1):
            p = doc.add_paragraph(style="List Number")
            p.add_run(s)

    # 5.2+ Scenarios
    for scn in scenarios:
        name = str(scn.get("name") or scn.get("id") or "Scenario")
        doc.add_heading(name, level=2)
        summary = str(scn.get("summary") or scn.get("description") or "").strip()
        if summary:
            doc.add_paragraph(summary)
        steps = scn.get("steps") if isinstance(scn.get("steps"), list) else []
        rows = []
        for st in steps:
            if not isinstance(st, dict):
                continue
            cycle = st.get("cycle")
            rows.append([
                _ssot_md_scalar(cycle if cycle is not None else ""),
                _ssot_md_scalar(st.get("action") or st.get("description") or ""),
                _ssot_md_scalar(st.get("fl_state") or st.get("function_state") or ""),
                _ssot_md_scalar(st.get("cl_state") or st.get("cycle_state") or st.get("stage") or ""),
            ])
        if rows:
            _ssot_docx_table_from_rows(doc, ["Cycle", "Action", "FL state", "CL state"], rows)

def _ssot_to_docx(data: dict, ip: str, out_path: Path) -> None:
    from docx import Document  # type: ignore

    doc = Document()
    safe_data = data if isinstance(data, dict) else {}
    _ssot_docx_set_footer(doc, ip)
    _ssot_docx_apply_heading_numbering(doc)
    _ssot_docx_cover_page(doc, ip, safe_data)
    _ssot_docx_render_revision_history(doc, safe_data)
    _ssot_docx_add_toc(doc)
    _ssot_docx_add_list_of_tables(doc)
    _ssot_docx_add_list_of_figures(doc)

    if not isinstance(data, dict):
        doc.add_heading("Raw Document", level=1)
        _ssot_docx_yaml_block(doc, data)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return

    # ── Chapter 1. Introduction ────────────────────────────────────
    doc.add_heading("Introduction", level=1)
    top = data.get("top_module") if isinstance(data.get("top_module"), dict) else {}
    intro_text = str(top.get("description") or top.get("purpose") or "").strip()
    if intro_text:
        doc.add_paragraph(intro_text)
    if not _ssot_section_is_empty(data.get("top_module")) or not _ssot_section_is_empty(data.get("sub_modules")):
        doc.add_heading("Block Diagram", level=2)
        _ssot_docx_render_block_diagram(doc, data)
    if not _ssot_section_is_empty(data.get("features")):
        doc.add_heading("Features", level=2)
        _ssot_docx_render_features(doc, data["features"])
    if (not _ssot_section_is_empty(data.get("features"))
            or not _ssot_section_is_empty(data.get("sub_modules"))
            or not _ssot_section_is_empty(data.get("function_model"))):
        doc.add_heading("Function Description", level=2)
        _ssot_docx_render_function_description(doc, data)

    # ── Chapter 2. Signal Description ──────────────────────────────
    if not _ssot_section_is_empty(data.get("io_list")):
        _ssot_docx_page_break(doc)
        doc.add_heading("Signal Description", level=1)
        _ssot_docx_render_io_list(doc, data["io_list"])

    # ── Chapter 3. Programming Model ───────────────────────────────
    if not _ssot_section_is_empty(data.get("registers")):
        _ssot_docx_page_break(doc)
        doc.add_heading("Programming Model", level=1)
        _ssot_docx_render_programming_model(doc, data["registers"])

    # ── Chapter 4. Hardware Configuration Options ──────────────────
    if not _ssot_section_is_empty(data.get("parameters")):
        _ssot_docx_page_break(doc)
        doc.add_heading("Hardware Configuration Options", level=1)
        _ssot_docx_render_parameters(doc, data["parameters"])

    # ── Chapter 5. Programming Sequence ────────────────────────────
    cycle_model = data.get("cycle_model") if isinstance(data.get("cycle_model"), dict) else {}
    fn_model = data.get("function_model") if isinstance(data.get("function_model"), dict) else {}
    has_scenarios = (isinstance(cycle_model.get("scenarios"), list) and cycle_model["scenarios"]) \
        or (isinstance(fn_model.get("scenarios"), list) and fn_model["scenarios"])
    if has_scenarios or not _ssot_section_is_empty(data.get("registers")):
        _ssot_docx_page_break(doc)
        doc.add_heading("Programming Sequence", level=1)
        _ssot_docx_render_programming_sequence(doc, data)

    # ── Appendix — Detailed sections ───────────────────────────────
    appendix_items = [
        (k, data.get(k)) for k in _SSOT_DOCX_APPENDIX_KEYS
        if k in data and not _ssot_section_is_empty(data.get(k))
    ]
    if appendix_items:
        _ssot_docx_page_break(doc)
        doc.add_heading("Appendix — Detailed Sections", level=1)
        for key, value in appendix_items:
            label = next((lbl for k, lbl in _SSOT_EXPORT_SECTION_ORDER if k == key), key.replace("_", " ").title())
            doc.add_heading(label, level=2)
            try:
                if key == "fsm":
                    _ssot_docx_render_fsm(doc, value)
                elif key == "interrupts":
                    _ssot_docx_render_interrupts(doc, value)
                elif key == "cycle_model":
                    _ssot_docx_render_cycle_model(doc, value)
                elif key == "error_handling":
                    _ssot_docx_render_error_handling(doc, value)
                else:
                    _ssot_docx_render_section(doc, key, value)
            except Exception as exc:
                err = doc.add_paragraph()
                err.add_run(f"(render error: {exc})").italic = True
                _ssot_docx_yaml_block(doc, value)

    # Catch-all: any keys not in the canonical list nor the appendix.
    canonical = {key for key, _ in _SSOT_EXPORT_SECTION_ORDER}
    chapter1_keys = {"top_module", "features", "sub_modules", "function_model"}
    chapter_keys = chapter1_keys | {"io_list", "registers", "parameters"} | set(_SSOT_DOCX_APPENDIX_KEYS)
    other_keys = [k for k in data.keys()
                  if k not in chapter_keys and k not in canonical
                  and not _ssot_section_is_empty(data.get(k))]
    if other_keys:
        _ssot_docx_page_break(doc)
        doc.add_heading("Other Sections", level=1)
        for key in other_keys:
            doc.add_heading(str(key), level=2)
            _ssot_docx_render_section(doc, key, data.get(key))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
