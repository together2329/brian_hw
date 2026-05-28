"""SSOT DOCX export — extracted from src/atlas_ui.py.

Renders the SSOT YAML to a Microsoft Word .docx datasheet (python-docx).
Companion to src/atlas_ssot_export.py (HTML/MD export) — same SSOT data
source, different output format.

Phase 2 of refactor/atlas-modular: move-only (no behavior change).
"""
from __future__ import annotations

import html as _html
import io
import json as _json
import os
import re
import re as _re
from pathlib import Path
from typing import Any, Optional

import yaml as _yaml

from src.atlas_ssot_export import (
    _ssot_md_scalar,
    _ssot_section_is_empty,
    _ssot_field_bit_info,
)


def _project_root() -> Path:
    """Read PROJECT_ROOT from atlas_ui dynamically (respects test monkeypatch)."""
    try:
        from src.atlas_ui import PROJECT_ROOT
        return PROJECT_ROOT
    except Exception:
        return Path(os.getcwd())


def _ssot_docx_set_mono(run: Any) -> None:
    try:
        run.font.name = "Consolas"
    except Exception:
        pass

def _ssot_docx_add_kv(doc: Any, key: str, value: Any) -> None:
    para = doc.add_paragraph()
    bold = para.add_run(f"{str(key).replace('_', ' ')}: ")
    bold.bold = True
    para.add_run(_ssot_md_scalar(value))

def _ssot_docx_yaml_block(doc: Any, value: Any) -> None:
    import yaml as _yaml  # type: ignore

    try:
        dumped = _yaml.safe_dump(value, allow_unicode=True, sort_keys=False).rstrip()
    except Exception as exc:
        dumped = f"(unable to render: {exc})"
    para = doc.add_paragraph()
    run = para.add_run(dumped)
    _ssot_docx_set_mono(run)

def _ssot_docx_list_of_dicts(doc: Any, rows: list) -> None:
    if not rows:
        return
    keys: list[str] = []
    seen: set[str] = set()
    for row in rows:
        if not isinstance(row, dict):
            continue
        for key in row.keys():
            if key not in seen:
                seen.add(key)
                keys.append(str(key))
    if not keys:
        _ssot_docx_yaml_block(doc, rows)
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(keys))
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        pass
    for col_idx, key in enumerate(keys):
        cell = table.rows[0].cells[col_idx]
        cell.text = key.replace("_", " ")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, key in enumerate(keys):
            raw = row.get(key) if isinstance(row, dict) else ""
            if isinstance(raw, (list, dict)):
                import yaml as _yaml  # type: ignore

                text = _yaml.safe_dump(raw, allow_unicode=True, sort_keys=False).strip()
            else:
                text = _ssot_md_scalar(raw)
            table.rows[row_idx].cells[col_idx].text = text

def _ssot_docx_dict_block(doc: Any, data: dict) -> None:
    if not isinstance(data, dict) or not data:
        return
    for key, value in data.items():
        if _ssot_md_is_short_scalar(value):
            _ssot_docx_add_kv(doc, str(key), value)
        elif isinstance(value, str):
            para = doc.add_paragraph()
            bold = para.add_run(f"{str(key).replace('_', ' ')}:")
            bold.bold = True
            doc.add_paragraph(value)
        elif isinstance(value, list) and value and all(isinstance(v, dict) for v in value):
            para = doc.add_paragraph()
            bold = para.add_run(f"{str(key).replace('_', ' ')}:")
            bold.bold = True
            _ssot_docx_list_of_dicts(doc, value)
        elif isinstance(value, list) and all(_ssot_md_is_short_scalar(v) for v in value):
            para = doc.add_paragraph()
            bold = para.add_run(f"{str(key).replace('_', ' ')}: ")
            bold.bold = True
            para.add_run(", ".join(_ssot_md_scalar(v) for v in value))
        elif isinstance(value, dict):
            para = doc.add_paragraph()
            bold = para.add_run(f"{str(key).replace('_', ' ')}:")
            bold.bold = True
            _ssot_docx_dict_block(doc, value)
        else:
            _ssot_docx_yaml_block(doc, value)

def _ssot_docx_render_section(doc: Any, key: str, value: Any) -> None:
    if key == "io_list" and isinstance(value, dict):
        cd = value.get("clock_domains")
        if cd:
            doc.add_heading("Clock Domains", level=2)
            if isinstance(cd, list) and cd and all(isinstance(r, dict) for r in cd):
                _ssot_docx_list_of_dicts(doc, cd)
            else:
                _ssot_docx_yaml_block(doc, cd)
        rst = value.get("resets")
        if rst:
            doc.add_heading("Resets", level=2)
            if isinstance(rst, list) and rst and all(isinstance(r, dict) for r in rst):
                _ssot_docx_list_of_dicts(doc, rst)
            else:
                _ssot_docx_yaml_block(doc, rst)
        ifs = value.get("interfaces")
        if isinstance(ifs, list):
            doc.add_heading("Interfaces", level=2)
            for iface in ifs:
                if not isinstance(iface, dict):
                    _ssot_docx_yaml_block(doc, iface)
                    continue
                doc.add_heading(str(iface.get("name") or "(unnamed)"), level=3)
                meta = {k: v for k, v in iface.items() if k != "ports"}
                _ssot_docx_dict_block(doc, meta)
                ports = iface.get("ports")
                if isinstance(ports, list) and ports:
                    _ssot_docx_list_of_dicts(doc, ports)
        leftover = {k: v for k, v in value.items()
                    if k not in ("clock_domains", "resets", "interfaces")}
        if leftover:
            _ssot_docx_dict_block(doc, leftover)
        return

    if key == "registers" and isinstance(value, dict):
        cfg = value.get("config")
        if isinstance(cfg, dict) and cfg:
            doc.add_heading("Config", level=2)
            _ssot_docx_dict_block(doc, cfg)
        reg_list = value.get("register_list")
        if isinstance(reg_list, list) and reg_list:
            doc.add_heading("Register List", level=2)
            _ssot_docx_list_of_dicts(doc, reg_list)
            for reg in reg_list:
                if not isinstance(reg, dict):
                    continue
                fields = reg.get("fields")
                if isinstance(fields, list) and fields:
                    doc.add_heading(
                        f"{reg.get('name') or '(unnamed reg)'} — fields",
                        level=3,
                    )
                    _ssot_docx_list_of_dicts(doc, fields)
        leftover = {k: v for k, v in value.items() if k not in ("config", "register_list")}
        if leftover:
            _ssot_docx_dict_block(doc, leftover)
        return

    if key == "features" and isinstance(value, list):
        for feat in value:
            if not isinstance(feat, dict):
                _ssot_docx_yaml_block(doc, feat)
                continue
            doc.add_heading(str(feat.get("name") or "(unnamed feature)"), level=2)
            rest = {k: v for k, v in feat.items() if k != "name"}
            _ssot_docx_dict_block(doc, rest)
        return

    if isinstance(value, dict):
        _ssot_docx_dict_block(doc, value)
        return
    if isinstance(value, list):
        if value and all(isinstance(r, dict) for r in value):
            _ssot_docx_list_of_dicts(doc, value)
            return
        if value and all(_ssot_md_is_short_scalar(r) for r in value):
            for item in value:
                doc.add_paragraph(_ssot_md_scalar(item), style="List Bullet")
            return
        _ssot_docx_yaml_block(doc, value)
        return
    if _ssot_md_is_short_scalar(value):
        doc.add_paragraph(_ssot_md_scalar(value))
        return
    _ssot_docx_yaml_block(doc, value)

def _ssot_docx_cover_page(doc: Any, ip: str, data: dict) -> None:
    """TRM-style cover page: product title, version, abstract, timestamp."""
    from datetime import datetime, timezone
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Pt  # type: ignore

    top = data.get("top_module") if isinstance(data, dict) else {}
    if not isinstance(top, dict):
        top = {}
    name = str(top.get("name") or ip).strip() or ip
    version = str(top.get("version") or "").strip() or "draft"
    description = str(top.get("description") or "").strip()
    kind = str(top.get("type") or "").strip()
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    blank = doc.add_paragraph()
    blank.paragraph_format.space_after = Pt(80)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(name)
    run.bold = True
    run.font.size = Pt(36)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run("Technical Reference Manual")
    sub.italic = True
    sub.font.size = Pt(18)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(60)

    if description:
        abstract_label = doc.add_paragraph()
        abstract_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl = abstract_label.add_run("Abstract")
        lbl.bold = True
        lbl.font.size = Pt(11)
        abstract = doc.add_paragraph(description)
        abstract.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in abstract.runs:
            run.font.size = Pt(11)

    spacer2 = doc.add_paragraph()
    spacer2.paragraph_format.space_after = Pt(40)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_lines = [
        f"Version: {version}",
        f"Module type: {kind}" if kind else "",
        f"Source: {ip}/yaml/{ip}.ssot.yaml",
        f"Exported: {ts}",
    ]
    for line in [l for l in meta_lines if l]:
        run = meta.add_run(line + "\n")
        run.font.size = Pt(10)
        run.italic = True

    _ssot_docx_page_break(doc)

def _ssot_docx_render_revision_history(doc: Any, data: dict) -> bool:
    """If top_module.history / data.history / data.revision_history exists, emit a TRM revision table."""
    if not isinstance(data, dict):
        return False
    history = None
    for src in (data.get("revision_history"), data.get("history"),
                (data.get("top_module") or {}).get("history") if isinstance(data.get("top_module"), dict) else None):
        if isinstance(src, list) and src:
            history = src
            break
    if not history:
        return False
    rows = []
    for ent in history:
        if isinstance(ent, dict):
            rows.append([
                _ssot_md_scalar(ent.get("version") or ent.get("rev") or ent.get("revision") or ""),
                _ssot_md_scalar(ent.get("date") or ent.get("when") or ""),
                _ssot_md_scalar(ent.get("author") or ent.get("by") or ""),
                _ssot_md_scalar(ent.get("changes") or ent.get("description") or ent.get("notes") or ""),
            ])
        else:
            rows.append(["", "", "", _ssot_md_scalar(ent)])
    if not rows:
        return False
    doc.add_heading("Revision History", level=1)
    _ssot_docx_table_from_rows(doc, ["Version", "Date", "Author", "Changes"], rows)
    _ssot_docx_page_break(doc)
    return True

def _ssot_docx_apply_heading_numbering(doc: Any) -> None:
    """Inject a multilevel-list style so Heading 1/2/3 auto-number as 1, 1.1, 1.1.1."""
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore

    numbering = doc.part.numbering_part.element
    # Use a high abstractNumId / numId to avoid colliding with anything Word
    # ships as default in this skeleton.
    abstract_id = "8801"
    num_id = "8801"
    # Skip if we already injected this.
    abstract_tag = qn("w:abstractNum")
    abstract_attr = qn("w:abstractNumId")
    for child in numbering:
        if child.tag == abstract_tag and child.get(abstract_attr) == abstract_id:
            return
    from docx.oxml import parse_xml  # type: ignore
    w = "xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'"
    abstract_xml = f"""
    <w:abstractNum {w} w:abstractNumId='{abstract_id}'>
      <w:multiLevelType w:val='hybridMultilevel'/>
      <w:lvl w:ilvl='0'>
        <w:start w:val='1'/>
        <w:numFmt w:val='decimal'/>
        <w:lvlText w:val='%1.'/>
        <w:lvlJc w:val='left'/>
        <w:pPr><w:ind w:left='432' w:hanging='432'/></w:pPr>
      </w:lvl>
      <w:lvl w:ilvl='1'>
        <w:start w:val='1'/>
        <w:numFmt w:val='decimal'/>
        <w:lvlText w:val='%1.%2'/>
        <w:lvlJc w:val='left'/>
        <w:pPr><w:ind w:left='720' w:hanging='576'/></w:pPr>
      </w:lvl>
      <w:lvl w:ilvl='2'>
        <w:start w:val='1'/>
        <w:numFmt w:val='decimal'/>
        <w:lvlText w:val='%1.%2.%3'/>
        <w:lvlJc w:val='left'/>
        <w:pPr><w:ind w:left='1008' w:hanging='720'/></w:pPr>
      </w:lvl>
    </w:abstractNum>"""
    numbering.append(parse_xml(abstract_xml))
    num_xml = f"""
    <w:num {w} w:numId='{num_id}'>
      <w:abstractNumId w:val='{abstract_id}'/>
    </w:num>"""
    numbering.append(parse_xml(num_xml))
    # Link Heading 1/2/3 styles to this numbering list.
    styles = doc.styles.element
    for level, style_name in enumerate(("Heading 1", "Heading 2", "Heading 3")):
        style = doc.styles[style_name]
        pPr = style.element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            style.element.append(pPr)
        # Drop any existing numPr in case we ran before.
        existing = pPr.find(qn("w:numPr"))
        if existing is not None:
            pPr.remove(existing)
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), str(level))
        numId = OxmlElement("w:numId"); numId.set(qn("w:val"), num_id)
        numPr.append(ilvl); numPr.append(numId)
        pPr.append(numPr)

def _ssot_docx_set_footer(doc: Any, ip: str) -> None:
    """TRM-style footer: '<ip> · <tab> · Page X of Y'."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.shared import Pt  # type: ignore

    for section in doc.sections:
        footer = section.footer
        # Replace any default paragraph so we don't double-render.
        if footer.paragraphs:
            para = footer.paragraphs[0]
            for run in list(para.runs):
                run.text = ""
        else:
            para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        left_run = para.add_run(f"{ip}  ·  ")
        left_run.font.size = Pt(9)
        left_run.italic = True

        page_label = para.add_run("Page ")
        page_label.font.size = Pt(9)

        # PAGE field
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "1"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        page_run = para.add_run()
        page_run.font.size = Pt(9)
        r = page_run._r
        r.append(fld_begin); r.append(instr); r.append(fld_sep); r.append(placeholder); r.append(fld_end)

        of_run = para.add_run(" of ")
        of_run.font.size = Pt(9)

        # NUMPAGES field
        fld_begin2 = OxmlElement("w:fldChar")
        fld_begin2.set(qn("w:fldCharType"), "begin")
        instr2 = OxmlElement("w:instrText")
        instr2.set(qn("xml:space"), "preserve")
        instr2.text = " NUMPAGES "
        fld_sep2 = OxmlElement("w:fldChar")
        fld_sep2.set(qn("w:fldCharType"), "separate")
        placeholder2 = OxmlElement("w:t")
        placeholder2.text = "1"
        fld_end2 = OxmlElement("w:fldChar")
        fld_end2.set(qn("w:fldCharType"), "end")
        num_run = para.add_run()
        num_run.font.size = Pt(9)
        r2 = num_run._r
        r2.append(fld_begin2); r2.append(instr2); r2.append(fld_sep2); r2.append(placeholder2); r2.append(fld_end2)

def _ssot_docx_page_break(doc: Any) -> None:
    from docx.enum.text import WD_BREAK  # type: ignore
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)

def _ssot_docx_add_word_index(doc: Any, heading: str, instr_text: str, placeholder: str = "") -> None:
    """Emit a Word field-based index block (TOC / List of Tables / List of Figures)."""
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore

    doc.add_heading(heading, level=1)
    para = doc.add_paragraph()
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder_el = OxmlElement("w:t")
    placeholder_el.text = placeholder or "Right-click → Update Field to populate."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder_el)
    r.append(fld_end)
    _ssot_docx_page_break(doc)

def _ssot_docx_add_toc(doc: Any) -> None:
    _ssot_docx_add_word_index(
        doc, "Table of Contents",
        r' TOC \o "1-3" \h \z \u ',
        "Right-click and choose 'Update Field' to populate the table of contents.",
    )

def _ssot_docx_add_list_of_tables(doc: Any) -> None:
    _ssot_docx_add_word_index(
        doc, "List of Tables",
        r' TOC \h \z \c "Table" ',
        "(Empty until you add a 'Table N: …' caption to each table and refresh fields.)",
    )

def _ssot_docx_add_list_of_figures(doc: Any) -> None:
    _ssot_docx_add_word_index(
        doc, "List of Figures",
        r' TOC \h \z \c "Figure" ',
        "(Empty until you add a 'Figure N: …' caption to each figure and refresh fields.)",
    )

def _ssot_docx_caption(doc: Any, kind: str, text: str) -> None:
    """Emit a 'Table N: text' or 'Figure N: text' caption using a SEQ counter."""
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Pt  # type: ignore

    para = doc.add_paragraph(style="Caption" if "Caption" in [s.name for s in doc.styles] else None)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = para.add_run(f"{kind} ")
    label_run.italic = True
    label_run.font.size = Pt(9)
    seq_run = para.add_run()
    seq_run.font.size = Pt(9)
    seq_run.italic = True
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' SEQ {kind} \\* ARABIC '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder_el = OxmlElement("w:t")
    placeholder_el.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    sr = seq_run._r
    sr.append(fld_begin); sr.append(instr); sr.append(fld_sep); sr.append(placeholder_el); sr.append(fld_end)
    tail_run = para.add_run(f"  {text}")
    tail_run.italic = True
    tail_run.font.size = Pt(9)

def _ssot_docx_table_from_rows(
    doc: Any,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    """TRM-style table with bold header row + light grid styling."""
    if not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        pass
    for col_idx, h in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = "" if value is None else str(value)

def _ssot_docx_render_top_module(doc: Any, top: dict) -> None:
    """Top-of-chapter prose + properties table instead of YAML dump."""
    if not isinstance(top, dict):
        _ssot_docx_yaml_block(doc, top)
        return
    description = str(top.get("description") or "").strip()
    purpose = str(top.get("purpose") or "").strip()
    if description:
        doc.add_paragraph(description)
    if purpose and purpose != description:
        doc.add_paragraph(purpose)
    prop_keys = [
        ("name", "Name"),
        ("file", "Top file"),
        ("type", "Module type"),
        ("version", "Version"),
        ("clock_freq_mhz", "Clock frequency (MHz)"),
        ("bus", "Bus"),
        ("role", "Role"),
    ]
    prop_rows = []
    for k, label in prop_keys:
        v = top.get(k)
        if v not in (None, "", [], {}):
            prop_rows.append([label, _ssot_md_scalar(v)])
    if prop_rows:
        doc.add_heading("Properties", level=2)
        _ssot_docx_table_from_rows(doc, ["Property", "Value"], prop_rows)
    nested_params = top.get("parameters")
    if nested_params:
        doc.add_heading("Top-Module Parameters", level=2)
        _ssot_docx_render_parameters(doc, nested_params)
    leftover = {k: v for k, v in top.items() if k not in dict(prop_keys) and k not in (
        "description", "purpose", "parameters", "history",
    )}
    if leftover:
        doc.add_heading("Additional Top-Module Fields", level=2)
        _ssot_docx_dict_block(doc, leftover)

def _ssot_docx_render_parameters(doc: Any, params: Any) -> None:
    """Name | Default | Type/Range | Description table."""
    rows_in: list[dict] = []
    if isinstance(params, list):
        for p in params:
            if isinstance(p, dict):
                rows_in.append(p)
    elif isinstance(params, dict):
        # KEY: VALUE mapping
        for k, v in params.items():
            if isinstance(v, dict):
                row = {"name": k, **v}
            else:
                row = {"name": k, "default": v}
            rows_in.append(row)
    if not rows_in:
        _ssot_docx_yaml_block(doc, params)
        return
    rows = []
    for r in rows_in:
        name = _ssot_md_scalar(r.get("name") or r.get("key") or r.get("param") or "")
        default = _ssot_md_scalar(
            r.get("default")
            or r.get("value")
            or r.get("default_value")
            or r.get("val")
            or ""
        )
        typ = _ssot_md_scalar(
            r.get("type") or r.get("range") or r.get("width") or ""
        )
        desc = _ssot_md_scalar(r.get("description") or r.get("desc") or r.get("notes") or "")
        rows.append([name, default, typ, desc])
    _ssot_docx_caption(doc, "Table", "Hardware configuration parameters")
    _ssot_docx_table_from_rows(doc, ["Parameter", "Default", "Type / Range", "Description"], rows)

def _ssot_docx_render_features(doc: Any, features: Any) -> None:
    """Numbered list with bold name + description paragraph."""
    if not isinstance(features, list):
        _ssot_docx_yaml_block(doc, features)
        return
    for idx, feat in enumerate(features, start=1):
        if not isinstance(feat, dict):
            doc.add_paragraph(_ssot_md_scalar(feat), style="List Number")
            continue
        name = str(feat.get("name") or f"Feature {idx}").strip() or f"Feature {idx}"
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(name)
        run.bold = True
        for field in ("description", "trigger", "datapath", "control", "output"):
            text = str(feat.get(field) or "").strip()
            if not text:
                continue
            sub = doc.add_paragraph()
            sub.paragraph_format.left_indent = sub.paragraph_format.left_indent or None
            label = sub.add_run(f"{field.replace('_', ' ').title()}: ")
            label.bold = True
            sub.add_run(text)

def _ssot_docx_render_register_detail(doc: Any, reg: dict, level: int = 2) -> None:
    """Per-register heading + properties table + bit-field table."""
    if not isinstance(reg, dict):
        _ssot_docx_yaml_block(doc, reg)
        return
    name = str(reg.get("name") or "(unnamed register)").strip()
    offset = _ssot_md_scalar(reg.get("offset") or "")
    title = f"{name}" + (f"  (@ {offset})" if offset else "")
    doc.add_heading(title, level=level)
    description = str(reg.get("description") or "").strip()
    if description:
        doc.add_paragraph(description)
    prop_keys = [
        ("offset", "Offset"),
        ("size", "Size"),
        ("width", "Width"),
        ("access", "Access"),
        ("reset", "Reset value"),
        ("type", "Type"),
    ]
    prop_rows = []
    for k, label in prop_keys:
        v = reg.get(k)
        if v not in (None, "", [], {}):
            prop_rows.append([label, _ssot_md_scalar(v)])
    if prop_rows:
        _ssot_docx_table_from_rows(doc, ["Property", "Value"], prop_rows)
    fields = reg.get("fields")
    if isinstance(fields, list) and fields:
        bit_rows = []
        for f in fields:
            if not isinstance(f, dict):
                continue
            bit_display, _lsb, width = _ssot_field_bit_info(f)
            bit_rows.append([
                bit_display,
                "" if width is None else str(width),
                _ssot_md_scalar(f.get("name") or ""),
                _ssot_md_scalar(f.get("access") or ""),
                _ssot_md_scalar(f.get("reset") or ""),
                _ssot_md_scalar(f.get("description") or ""),
            ])
        if bit_rows:
            field_label = doc.add_paragraph()
            field_label.add_run("Bit fields").bold = True
            _ssot_docx_table_from_rows(
                doc,
                ["Bits", "Width", "Field", "Access", "Reset", "Description"],
                bit_rows,
            )

def _ssot_docx_port_direction(port: dict) -> str:
    raw = str(port.get("direction") or port.get("dir") or port.get("io") or "").strip().lower()
    if not raw:
        return "Unspecified"
    if raw in {"in", "input", "i"}:
        return "Input"
    if raw in {"out", "output", "o"}:
        return "Output"
    if raw in {"inout", "bidir", "io", "b"}:
        return "Bidirectional"
    return raw.title()

def _ssot_docx_port_rows(ports: list, include_iface: bool = False, iface_name: str = "") -> dict[str, list[list[str]]]:
    """Bucket ports by canonical direction → list of [Signal, Width, Clock, Reset, Description (, Iface)]."""
    bucketed: dict[str, list[list[str]]] = {}
    for port in ports:
        if not isinstance(port, dict):
            continue
        direction = _ssot_docx_port_direction(port)
        row = [
            _ssot_md_scalar(port.get("name") or port.get("signal") or ""),
            _ssot_md_scalar(port.get("width") or port.get("bits") or port.get("range") or ""),
            _ssot_md_scalar(port.get("clock") or port.get("clock_domain") or ""),
            _ssot_md_scalar(port.get("reset") or port.get("reset_domain") or ""),
            _ssot_md_scalar(port.get("description") or port.get("desc") or ""),
        ]
        if include_iface:
            row.append(iface_name)
        bucketed.setdefault(direction, []).append(row)
    return bucketed

def _ssot_docx_render_port_buckets(
    doc: Any,
    bucketed: dict[str, list[list[str]]],
    include_iface: bool = False,
) -> None:
    if not bucketed:
        return
    headers = ["Signal", "Width", "Clock", "Reset", "Description"]
    if include_iface:
        headers.append("Interface")
    ordering = ["Input", "Output", "Bidirectional", "Unspecified"]
    seen = set()
    keys = [k for k in ordering if k in bucketed] + [k for k in bucketed if k not in ordering]
    for direction in keys:
        if direction in seen:
            continue
        seen.add(direction)
        rows = bucketed.get(direction, [])
        if not rows:
            continue
        sub = doc.add_paragraph()
        sub.add_run(f"{direction} signals  ({len(rows)})").bold = True
        _ssot_docx_table_from_rows(doc, headers, rows)

def _ssot_docx_render_io_list(doc: Any, value: Any) -> None:
    """Clock/Reset/Interface (with direction-grouped port tables) signal description."""
    if isinstance(value, list):
        bucketed = _ssot_docx_port_rows(value)
        if bucketed:
            _ssot_docx_render_port_buckets(doc, bucketed)
            return
        _ssot_docx_yaml_block(doc, value)
        return

    if not isinstance(value, dict):
        _ssot_docx_yaml_block(doc, value)
        return

    cd = value.get("clock_domains") or value.get("clocks")
    if isinstance(cd, list) and cd:
        doc.add_heading("Clock domains", level=2)
        rows = []
        for entry in cd:
            if not isinstance(entry, dict):
                continue
            rows.append([
                _ssot_md_scalar(entry.get("name") or ""),
                _ssot_md_scalar(entry.get("frequency_mhz") or entry.get("frequency") or ""),
                _ssot_md_scalar(entry.get("source") or ""),
                _ssot_md_scalar(entry.get("description") or ""),
            ])
        if rows:
            _ssot_docx_caption(doc, "Table", "Clock domains")
            _ssot_docx_table_from_rows(doc, ["Domain", "Frequency (MHz)", "Source", "Description"], rows)

    rst = value.get("resets")
    if isinstance(rst, list) and rst:
        doc.add_heading("Resets", level=2)
        rows = []
        for entry in rst:
            if not isinstance(entry, dict):
                continue
            rows.append([
                _ssot_md_scalar(entry.get("name") or ""),
                _ssot_md_scalar(entry.get("polarity") or ""),
                _ssot_md_scalar(entry.get("sync_async") or entry.get("type") or ""),
                _ssot_md_scalar(entry.get("source") or ""),
                _ssot_md_scalar(entry.get("description") or ""),
            ])
        if rows:
            _ssot_docx_caption(doc, "Table", "Reset sources")
            _ssot_docx_table_from_rows(doc, ["Reset", "Polarity", "Sync / Async", "Source", "Description"], rows)

    ifs = value.get("interfaces")
    if isinstance(ifs, list) and ifs:
        doc.add_heading("Interfaces", level=2)
        for iface in ifs:
            if not isinstance(iface, dict):
                _ssot_docx_yaml_block(doc, iface)
                continue
            iname = str(iface.get("name") or "(unnamed)")
            doc.add_heading(iname, level=3)
            description = str(iface.get("description") or "").strip()
            if description:
                doc.add_paragraph(description)
            iface_meta_keys = [
                ("type", "Type"),
                ("role", "Role"),
                ("protocol", "Protocol"),
                ("data_width", "Data width"),
                ("addr_width", "Address width"),
                ("clock", "Clock domain"),
                ("reset", "Reset domain"),
            ]
            meta_rows = []
            for k, label in iface_meta_keys:
                v = iface.get(k)
                if v not in (None, "", [], {}):
                    meta_rows.append([label, _ssot_md_scalar(v)])
            if meta_rows:
                _ssot_docx_table_from_rows(doc, ["Property", "Value"], meta_rows)
            ports = iface.get("ports")
            if isinstance(ports, list) and ports:
                bucketed = _ssot_docx_port_rows(ports)
                _ssot_docx_caption(doc, "Table", f"Signal description — {iname}")
                _ssot_docx_render_port_buckets(doc, bucketed)

    flat_ports = value.get("ports")
    if isinstance(flat_ports, list) and flat_ports and not ifs:
        doc.add_heading("Signals", level=2)
        bucketed = _ssot_docx_port_rows(flat_ports)
        _ssot_docx_render_port_buckets(doc, bucketed)

    leftover = {
        k: v for k, v in value.items()
        if k not in ("clock_domains", "clocks", "resets", "interfaces", "ports")
        and not _ssot_section_is_empty(v)
    }
    if leftover:
        doc.add_heading("Additional I/O metadata", level=2)
        _ssot_docx_dict_block(doc, leftover)

def _ssot_docx_render_interrupts(doc: Any, value: Any) -> None:
    """ID | Source | Polarity | Mask register | Description."""
    if isinstance(value, dict):
        config = {k: v for k, v in value.items() if k not in ("interrupt_list", "interrupts", "list")}
        items = value.get("interrupt_list") or value.get("interrupts") or value.get("list")
        if config:
            cfg_rows = [[k.replace("_", " "), _ssot_md_scalar(v)]
                        for k, v in config.items() if not _ssot_section_is_empty(v)]
            if cfg_rows:
                doc.add_heading("Interrupt configuration", level=2)
                _ssot_docx_table_from_rows(doc, ["Property", "Value"], cfg_rows)
    else:
        items = value
    if not isinstance(items, list) or not items:
        _ssot_docx_yaml_block(doc, value)
        return
    rows = []
    detail_blocks = []
    for idx, ent in enumerate(items, start=1):
        if not isinstance(ent, dict):
            rows.append(["", _ssot_md_scalar(ent), "", "", ""])
            continue
        irq_id = _ssot_md_scalar(ent.get("id") or ent.get("number") or ent.get("vector") or idx)
        name = _ssot_md_scalar(ent.get("name") or ent.get("signal") or "")
        source = _ssot_md_scalar(ent.get("source") or ent.get("cause") or "")
        polarity = _ssot_md_scalar(ent.get("polarity") or ent.get("level") or ent.get("trigger") or "")
        mask = _ssot_md_scalar(ent.get("mask") or ent.get("mask_register") or ent.get("enable") or "")
        description = _ssot_md_scalar(ent.get("description") or ent.get("desc") or "")
        rows.append([
            irq_id,
            (name + (f" ({source})" if source and source != name else "")) if name else source,
            polarity,
            mask,
            description[:120],
        ])
        detail_blocks.append((name or irq_id, ent, description))
    doc.add_heading("Interrupt summary", level=2)
    _ssot_docx_table_from_rows(
        doc,
        ["ID", "Signal / Source", "Polarity / Trigger", "Mask register", "Description"],
        rows,
    )
    # Detail blocks for any interrupts that carry richer metadata (clearing rule, status bit, etc).
    extra_keys = ("clear", "status", "ack", "priority", "vector_address", "context", "notes")
    detail_rows = []
    for label, ent, desc in detail_blocks:
        meta = []
        for k in extra_keys:
            if ent.get(k) not in (None, "", [], {}):
                meta.append((k, ent.get(k)))
        if meta:
            detail_rows.append((label, desc, meta))
    if detail_rows:
        doc.add_heading("Interrupt details", level=2)
        for label, desc, meta in detail_rows:
            doc.add_heading(str(label), level=3)
            if desc:
                doc.add_paragraph(desc)
            _ssot_docx_table_from_rows(
                doc,
                ["Property", "Value"],
                [[k.replace("_", " "), _ssot_md_scalar(v)] for k, v in meta],
            )

def _ssot_docx_render_fsm(doc: Any, value: Any) -> None:
    """State list + transition table per machine."""
    machines: list[dict] = []
    if isinstance(value, dict):
        ml = value.get("machines") or value.get("fsm_list") or value.get("list")
        if isinstance(ml, list):
            machines = [m for m in ml if isinstance(m, dict)]
        else:
            # Pattern: {<machine_name>: {states: [...], transitions: [...]}}
            for k, v in value.items():
                if isinstance(v, dict) and (v.get("states") or v.get("transitions")):
                    machines.append({"name": k, **v})
    elif isinstance(value, list):
        machines = [m for m in value if isinstance(m, dict)]
    if not machines:
        _ssot_docx_yaml_block(doc, value)
        return
    for m in machines:
        name = str(m.get("name") or "(unnamed FSM)")
        doc.add_heading(name, level=2)
        description = str(m.get("description") or m.get("purpose") or "").strip()
        if description:
            doc.add_paragraph(description)
        meta_keys = [
            ("reset_state", "Reset state"),
            ("encoding", "Encoding"),
            ("clock", "Clock domain"),
            ("reset", "Reset domain"),
            ("illegal_recovery", "Illegal recovery"),
        ]
        meta_rows = []
        for k, label in meta_keys:
            v = m.get(k)
            if v not in (None, "", [], {}):
                meta_rows.append([label, _ssot_md_scalar(v)])
        if meta_rows:
            _ssot_docx_table_from_rows(doc, ["Property", "Value"], meta_rows)
        states = m.get("states")
        if isinstance(states, list) and states:
            doc.add_heading("States", level=3)
            rows = []
            for s in states:
                if isinstance(s, dict):
                    rows.append([
                        _ssot_md_scalar(s.get("id") or s.get("name") or ""),
                        _ssot_md_scalar(s.get("label") or s.get("name") or ""),
                        _ssot_md_scalar(s.get("encoding") or ""),
                        "yes" if s.get("reset") else "",
                        _ssot_md_scalar(s.get("description") or s.get("action") or ""),
                    ])
                else:
                    rows.append(["", _ssot_md_scalar(s), "", "", ""])
            _ssot_docx_table_from_rows(doc, ["ID", "State", "Encoding", "Reset", "Description / Action"], rows)
        transitions = m.get("transitions")
        if isinstance(transitions, list) and transitions:
            doc.add_heading("Transitions", level=3)
            rows = []
            for tr in transitions:
                if isinstance(tr, dict):
                    rows.append([
                        _ssot_md_scalar(tr.get("from") or tr.get("source") or tr.get("current") or ""),
                        _ssot_md_scalar(tr.get("event") or tr.get("condition") or tr.get("trigger") or ""),
                        _ssot_md_scalar(tr.get("to") or tr.get("dest") or tr.get("target") or tr.get("next") or ""),
                        _ssot_md_scalar(tr.get("action") or tr.get("output") or ""),
                        _ssot_md_scalar(tr.get("description") or ""),
                    ])
                else:
                    rows.append(["", _ssot_md_scalar(tr), "", "", ""])
            _ssot_docx_table_from_rows(
                doc, ["Current", "Event / Condition", "Next", "Action / Output", "Description"], rows,
            )
        outputs = m.get("outputs") or m.get("output_rules")
        if outputs:
            doc.add_heading("Output rules", level=3)
            if isinstance(outputs, list) and outputs and all(isinstance(o, dict) for o in outputs):
                _ssot_docx_list_of_dicts(doc, outputs)
            else:
                _ssot_docx_yaml_block(doc, outputs)

def _ssot_docx_render_cycle_model(doc: Any, value: Any) -> None:
    """Cycle contract prose + Latency table + Handshake table + Pipeline table."""
    if not isinstance(value, dict):
        _ssot_docx_yaml_block(doc, value)
        return
    purpose = str(value.get("purpose") or value.get("description") or "").strip()
    if purpose:
        doc.add_paragraph(purpose)
    contract_keys = [
        ("clock", "Clock domain"),
        ("reset", "Reset domain"),
        ("assertion", "Reset assertion"),
        ("deassertion", "Reset deassertion"),
        ("scope", "Scope"),
    ]
    contract_rows = []
    for k, label in contract_keys:
        v = value.get(k)
        if v not in (None, "", [], {}):
            contract_rows.append([label, _ssot_md_scalar(v)])
    if contract_rows:
        doc.add_heading("Cycle contract", level=2)
        _ssot_docx_table_from_rows(doc, ["Property", "Value"], contract_rows)
    latency = value.get("latency") or value.get("latencies")
    if isinstance(latency, dict) and latency:
        doc.add_heading("Latency", level=2)
        rows = []
        for path, info in latency.items():
            if isinstance(info, dict):
                rows.append([
                    str(path),
                    _ssot_md_scalar(info.get("min_cycles") or info.get("min") or ""),
                    _ssot_md_scalar(info.get("max_cycles") or info.get("max") or ""),
                    _ssot_md_scalar(info.get("description") or info.get("notes") or ""),
                ])
            else:
                rows.append([str(path), "", "", _ssot_md_scalar(info)])
        _ssot_docx_table_from_rows(doc, ["Path", "Min cycles", "Max cycles", "Description"], rows)
    elif isinstance(latency, list) and latency:
        doc.add_heading("Latency", level=2)
        rows = []
        for ent in latency:
            if isinstance(ent, dict):
                rows.append([
                    _ssot_md_scalar(ent.get("path") or ent.get("name") or ent.get("from_to") or ""),
                    _ssot_md_scalar(ent.get("min_cycles") or ent.get("min") or ""),
                    _ssot_md_scalar(ent.get("max_cycles") or ent.get("max") or ""),
                    _ssot_md_scalar(ent.get("description") or ent.get("notes") or ""),
                ])
        _ssot_docx_table_from_rows(doc, ["Path", "Min cycles", "Max cycles", "Description"], rows)
    handshake = value.get("handshake_rules") or value.get("handshake")
    if isinstance(handshake, list) and handshake:
        doc.add_heading("Handshake rules", level=2)
        rows = []
        for ent in handshake:
            if isinstance(ent, dict):
                rows.append([
                    _ssot_md_scalar(ent.get("signal") or ent.get("name") or ""),
                    _ssot_md_scalar(ent.get("rule") or ent.get("contract") or ent.get("description") or ""),
                    _ssot_md_scalar(ent.get("notes") or ""),
                ])
        _ssot_docx_table_from_rows(doc, ["Signal", "Rule", "Notes"], rows)
    pipeline = value.get("pipeline") or value.get("stages")
    if isinstance(pipeline, list) and pipeline:
        doc.add_heading("Pipeline", level=2)
        rows = []
        for ent in pipeline:
            if isinstance(ent, dict):
                rows.append([
                    _ssot_md_scalar(ent.get("stage") or ent.get("name") or ent.get("id") or ""),
                    _ssot_md_scalar(ent.get("cycle") or ent.get("phase") or ""),
                    _ssot_md_scalar(ent.get("action") or ent.get("description") or ""),
                    _ssot_md_scalar(ent.get("ready") or ""),
                    _ssot_md_scalar(ent.get("notes") or ""),
                ])
        _ssot_docx_table_from_rows(doc, ["Stage", "Cycle", "Action", "Ready / Backpressure", "Notes"], rows)
    leftover = {
        k: v for k, v in value.items()
        if k not in dict(contract_keys)
        and k not in ("latency", "latencies", "handshake_rules", "handshake", "pipeline", "stages", "purpose", "description")
        and not _ssot_section_is_empty(v)
    }
    if leftover:
        doc.add_heading("Additional cycle-model fields", level=2)
        _ssot_docx_dict_block(doc, leftover)

def _ssot_docx_render_error_handling(doc: Any, value: Any) -> None:
    """Error-code table + per-error detail blocks."""
    errors: list = []
    leftover: dict = {}
    if isinstance(value, dict):
        for cand_key in ("errors", "error_list", "error_sources", "codes", "list"):
            cand = value.get(cand_key)
            if isinstance(cand, list) and cand:
                errors = [e for e in cand if isinstance(e, dict)]
                break
        leftover = {
            k: v for k, v in value.items()
            if k not in ("errors", "error_list", "error_sources", "codes", "list")
            and not _ssot_section_is_empty(v)
        }
        policy = str(value.get("policy") or value.get("strategy") or "").strip()
        if policy:
            doc.add_paragraph(policy)
    elif isinstance(value, list):
        errors = [e for e in value if isinstance(e, dict)]
    if not errors:
        if leftover:
            _ssot_docx_dict_block(doc, leftover)
        if not leftover:
            _ssot_docx_yaml_block(doc, value)
        return
    summary_rows = []
    for idx, e in enumerate(errors, start=1):
        summary_rows.append([
            _ssot_md_scalar(e.get("id") or e.get("code") or idx),
            _ssot_md_scalar(e.get("name") or e.get("label") or ""),
            _ssot_md_scalar(e.get("condition") or e.get("cause") or e.get("trigger") or ""),
            _ssot_md_scalar(e.get("severity") or ""),
            _ssot_md_scalar(e.get("recovery") or e.get("response") or e.get("handling") or ""),
        ])
    doc.add_heading("Error summary", level=2)
    _ssot_docx_table_from_rows(
        doc,
        ["ID / Code", "Name", "Condition / Cause", "Severity", "Recovery / Handling"],
        summary_rows,
    )
    detail_keys = ("architectural_effect", "status_register", "clear", "log",
                   "notes", "description", "isolation", "trace")
    detail_payload = [(e, [(k, e.get(k)) for k in detail_keys if e.get(k) not in (None, "", [], {})])
                      for e in errors]
    detail_payload = [(e, kvs) for e, kvs in detail_payload if kvs]
    if detail_payload:
        doc.add_heading("Error details", level=2)
        for e, kvs in detail_payload:
            label = str(e.get("name") or e.get("id") or e.get("code") or "(error)")
            doc.add_heading(label, level=3)
            _ssot_docx_table_from_rows(
                doc, ["Property", "Value"],
                [[k.replace("_", " "), _ssot_md_scalar(v)] for k, v in kvs],
            )
    if leftover:
        doc.add_heading("Additional error-handling metadata", level=2)
        _ssot_docx_dict_block(doc, leftover)

def _ssot_docx_block_diagram_png(data: dict):
    """Render an Andes-style block diagram PNG (BytesIO) via matplotlib.

    Returns None when matplotlib is not importable; the caller then
    emits only the text-and-table representation.
    """
    try:
        import matplotlib  # type: ignore
        matplotlib.use("Agg", force=True)
        import matplotlib.pyplot as plt  # type: ignore
        import matplotlib.patches as mpatches  # type: ignore
    except Exception:
        return None
    import io as _io

    top = data.get("top_module") if isinstance(data.get("top_module"), dict) else {}
    submods_raw = data.get("sub_modules") if isinstance(data.get("sub_modules"), list) else []
    submods = [sm for sm in submods_raw if isinstance(sm, (dict, str))]
    io_list = data.get("io_list") if isinstance(data.get("io_list"), dict) else {}
    interfaces = io_list.get("interfaces") if isinstance(io_list, dict) else None
    if not isinstance(interfaces, list):
        interfaces = []

    fig_w = 9.5
    fig_h = 6.0 if submods else 4.0
    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, fig_h)
    ax.axis("off")

    outer_x, outer_w = 1.5, 7.0
    outer_y, outer_h = 0.6, fig_h - 1.0
    outer = mpatches.FancyBboxPatch(
        (outer_x, outer_y), outer_w, outer_h,
        boxstyle="round,pad=0.02",
        edgecolor="#4a6cf7", facecolor="#f4f7ff", lw=2,
    )
    ax.add_patch(outer)
    top_name = str(top.get("name") or "TOP")
    top_role = str(top.get("type") or top.get("role") or "").strip()
    ax.text(
        outer_x + outer_w / 2, outer_y + outer_h - 0.35,
        top_name + (f"   ({top_role})" if top_role else ""),
        ha="center", va="center", fontsize=14, fontweight="bold",
    )

    if submods:
        cols = min(3, len(submods))
        rows = (len(submods) + cols - 1) // cols
        inner_w = outer_w - 0.6
        inner_h = outer_h - 1.4
        cell_w = inner_w / cols
        cell_h = inner_h / rows
        for i, sm in enumerate(submods):
            r = i // cols
            c = i % cols
            x = outer_x + 0.3 + c * cell_w
            y = outer_y + 0.4 + (rows - 1 - r) * cell_h
            sub = mpatches.FancyBboxPatch(
                (x + 0.08, y + 0.08), cell_w - 0.16, cell_h - 0.16,
                boxstyle="round,pad=0.02",
                edgecolor="#4a6cf7", facecolor="#ffffff", lw=1.2,
            )
            ax.add_patch(sub)
            name = (sm.get("name") if isinstance(sm, dict) else str(sm)) or "?"
            desc = ""
            if isinstance(sm, dict):
                desc = str(sm.get("description") or sm.get("implementation") or "")[:64]
            ax.text(
                x + cell_w / 2, y + cell_h / 2 + (0.15 if desc else 0),
                name, ha="center", va="center",
                fontsize=10, fontweight="bold", family="monospace",
            )
            if desc:
                ax.text(
                    x + cell_w / 2, y + cell_h / 2 - 0.2,
                    desc, ha="center", va="center",
                    fontsize=7, color="#5a6a8f",
                )

    valid_ifs = [i for i in interfaces if isinstance(i, dict)][:4]
    in_color, out_color = "#4a6cf7", "#7a4af7"
    for idx, ifc in enumerate(valid_ifs):
        name = str(ifc.get("name") or "?")
        ifc_type = str(ifc.get("type") or "").strip()
        role = str(ifc.get("role") or "").lower()
        inward = role in {"slave", "sink", "in", "input", ""}
        side_left = idx % 2 == 0
        y_pos = outer_y + outer_h - 0.9 - (idx // 2) * 1.6
        if y_pos < outer_y + 0.4:
            y_pos = outer_y + outer_h / 2
        if side_left:
            arrow = mpatches.FancyArrowPatch(
                (0.4, y_pos) if inward else (outer_x + 0.1, y_pos),
                (outer_x + 0.1, y_pos) if inward else (0.4, y_pos),
                arrowstyle="->", mutation_scale=18,
                color=in_color if inward else out_color, lw=1.6,
            )
            ax.text(0.55, y_pos + 0.22, name + (f"\n({ifc_type})" if ifc_type else ""),
                    ha="center", va="bottom", fontsize=9, color="#1f2a44")
        else:
            arrow = mpatches.FancyArrowPatch(
                (9.6, y_pos) if inward else (outer_x + outer_w - 0.1, y_pos),
                (outer_x + outer_w - 0.1, y_pos) if inward else (9.6, y_pos),
                arrowstyle="->", mutation_scale=18,
                color=in_color if inward else out_color, lw=1.6,
            )
            ax.text(9.45, y_pos + 0.22, name + (f"\n({ifc_type})" if ifc_type else ""),
                    ha="center", va="bottom", fontsize=9, color="#1f2a44")
        ax.add_patch(arrow)

    buf = _io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=160, facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf

def _ssot_docx_render_block_diagram(doc: Any, data: dict) -> None:
    """Block-diagram chapter body — real figure first, submodule table beneath as key."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Inches, Pt  # type: ignore

    fig_buf = _ssot_docx_block_diagram_png(data)
    if fig_buf is not None:
        doc.add_picture(fig_buf, width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _ssot_docx_caption(doc, "Figure", "Block diagram")

    top = data.get("top_module") if isinstance(data.get("top_module"), dict) else {}
    top_name = str(top.get("name") or "TOP").strip() or "TOP"
    top_role = str(top.get("type") or top.get("role") or "").strip()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(top_name + (f"  ({top_role})" if top_role else ""))
    title_run.bold = True
    title_run.font.size = Pt(13)

    submods = data.get("sub_modules")
    if isinstance(submods, list) and submods:
        cols = min(3, len(submods))
        rows = (len(submods) + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        try:
            table.style = "Light Grid Accent 1"
        except Exception:
            pass
        for idx, sm in enumerate(submods):
            r = idx // cols
            c = idx % cols
            cell = table.rows[r].cells[c]
            if isinstance(sm, dict):
                name = str(sm.get("name") or sm.get("module") or "?").strip()
                desc = str(sm.get("description") or sm.get("implementation") or "").strip()
            else:
                name = str(sm)
                desc = ""
            cell.text = name
            for run in cell.paragraphs[0].runs:
                run.bold = True
            if desc:
                p = cell.add_paragraph()
                p.add_run(desc[:160]).italic = True
        # Center each cell
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        caption = doc.add_paragraph()
        cap = caption.add_run("(no sub-modules declared — single-block design)")
        cap.italic = True
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # External interface chips
    io_list = data.get("io_list") if isinstance(data.get("io_list"), dict) else {}
    interfaces = io_list.get("interfaces") if isinstance(io_list, dict) else None
    if isinstance(interfaces, list) and interfaces:
        legend = doc.add_paragraph()
        legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
        legend.add_run("External interfaces: ").italic = True
        legend.add_run(", ".join(
            str(i.get("name") or "?") + (f" ({i.get('type')})" if isinstance(i, dict) and i.get("type") else "")
            for i in interfaces if isinstance(i, dict)
        ))

def _ssot_docx_render_function_description(doc: Any, data: dict) -> None:
    """Per-feature or per-submodule prose sub-section.

    Mirrors AndeShape §1.3 layout — one H3 per feature with the
    description paragraph that walks the reader through what that
    feature does. Falls through to sub_modules → function_model
    transactions when features are absent.
    """
    features = data.get("features")
    if isinstance(features, list) and features:
        for feat in features:
            if not isinstance(feat, dict):
                doc.add_paragraph(_ssot_md_scalar(feat))
                continue
            name = str(feat.get("name") or "(unnamed)")
            doc.add_heading(name, level=3)
            for body_key in ("description", "datapath", "control", "output", "trigger"):
                txt = str(feat.get(body_key) or "").strip()
                if txt:
                    p = doc.add_paragraph()
                    if body_key != "description":
                        p.add_run(f"{body_key.title()}: ").bold = True
                    p.add_run(txt)
        return
    submods = data.get("sub_modules")
    if isinstance(submods, list) and submods:
        for sm in submods:
            if not isinstance(sm, dict):
                continue
            name = str(sm.get("name") or sm.get("module") or "(unnamed)")
            desc = str(sm.get("description") or sm.get("implementation") or "").strip()
            if not desc:
                continue
            doc.add_heading(name, level=3)
            doc.add_paragraph(desc)
        return
    fn_model = data.get("function_model") if isinstance(data.get("function_model"), dict) else {}
    purpose = str(fn_model.get("purpose") or fn_model.get("description") or "").strip()
    if purpose:
        doc.add_paragraph(purpose)
    transactions = fn_model.get("transactions") if isinstance(fn_model, dict) else None
    if isinstance(transactions, list) and transactions:
        for tx in transactions:
            if not isinstance(tx, dict):
                continue
            label = str(tx.get("id") or tx.get("name") or "(transaction)").strip()
            doc.add_heading(label, level=3)
            for body_key in ("description", "preconditions", "inputs", "outputs", "side_effects"):
                v = tx.get(body_key)
                if v in (None, "", [], {}):
                    continue
                if isinstance(v, (list, tuple)):
                    text = "; ".join(_ssot_md_scalar(item) for item in v)
                else:
                    text = _ssot_md_scalar(v)
                p = doc.add_paragraph()
                p.add_run(f"{body_key.replace('_', ' ').title()}: ").bold = True
                p.add_run(text)

def _ssot_docx_render_programming_model(doc: Any, registers: Any) -> None:
    """Andes §3 layout — Summary of Registers table → per-register Description."""
    reg_list = []
    if isinstance(registers, dict):
        cfg = registers.get("config")
        if isinstance(cfg, dict) and cfg:
            doc.add_heading("Register Block Configuration", level=2)
            _ssot_docx_table_from_rows(
                doc, ["Property", "Value"],
                [[k.replace("_", " "), _ssot_md_scalar(v)] for k, v in cfg.items()],
            )
        if isinstance(registers.get("register_list"), list):
            reg_list = registers["register_list"]
    elif isinstance(registers, list):
        reg_list = registers
    reg_list = [r for r in reg_list if isinstance(r, dict)]
    if not reg_list:
        _ssot_docx_yaml_block(doc, registers)
        return
    doc.add_heading("Summary of Registers", level=2)
    _ssot_docx_caption(doc, "Table", "Summary of registers")
    summary_rows = []
    for reg in reg_list:
        summary_rows.append([
            _ssot_md_scalar(reg.get("offset") or ""),
            _ssot_md_scalar(reg.get("name") or ""),
            _ssot_md_scalar(reg.get("access") or ""),
            _ssot_md_scalar(reg.get("reset") or ""),
            _ssot_md_scalar(reg.get("description") or "")[:120],
        ])
    _ssot_docx_table_from_rows(
        doc, ["Offset", "Register", "Access", "Reset", "Description"], summary_rows,
    )
    doc.add_heading("Register Description", level=2)
    for reg in reg_list:
        _ssot_docx_render_register_detail(doc, reg, level=3)

def _ssot_docx_render_programming_sequence(doc: Any, data: dict) -> None:
    """Andes §5 layout — Setup / TX / RX style sub-sections derived from SSOT.

    Heuristics:
      * "IP Setup": find Enable / Configure registers, emit a numbered
        step list with the relevant register writes.
      * Scenarios: if cycle_model.scenarios / function_model.scenarios
        exist, one H3 per scenario walking cycle-by-cycle through the
        steps with action prose.
    """
    cycle_model = data.get("cycle_model") if isinstance(data.get("cycle_model"), dict) else {}
    fn_model = data.get("function_model") if isinstance(data.get("function_model"), dict) else {}
    scenarios = []
    for src in (cycle_model.get("scenarios"), fn_model.get("scenarios")):
        if isinstance(src, list):
            scenarios.extend(s for s in src if isinstance(s, dict))

    # 5.1 IP Setup
    registers = data.get("registers")
    reg_list = []
    if isinstance(registers, dict) and isinstance(registers.get("register_list"), list):
        reg_list = [r for r in registers["register_list"] if isinstance(r, dict)]
    elif isinstance(registers, list):
        reg_list = [r for r in registers if isinstance(r, dict)]
    enable_targets = []
    configure_targets = []
    for reg in reg_list:
        rn = str(reg.get("name") or "")
        fields = reg.get("fields") if isinstance(reg.get("fields"), list) else []
        en_field = next((f for f in fields if isinstance(f, dict)
                         and re.search(r"\ben\b|enable", str(f.get("name") or ""), re.I)), None)
        if en_field:
            enable_targets.append((rn, en_field.get("name"), en_field.get("description")))
        elif re.match(r"^(cr|ctrl|control|cfg|config|mode)\b", rn, re.I):
            enable_targets.append((rn, None, reg.get("description")))
        if re.match(r"^(dlr|brr|baud|div|mr|prescaler|cfg)", rn, re.I) and not any(rn == e[0] for e in enable_targets):
            configure_targets.append((rn, reg.get("offset"), reg.get("description")))
    if enable_targets or configure_targets:
        doc.add_heading("IP Setup", level=2)
        steps = []
        for rn, fld, desc in enable_targets[:2]:
            steps.append(f"Write {rn}.{fld}=1 to enable the module." if fld else f"Configure {rn} (the primary control register) to bring the module out of reset.")
        for rn, off, desc in configure_targets[:3]:
            label = f"{rn}" + (f" (@ {off})" if off else "")
            steps.append(f"Program {label} with the operating parameters required for your use case.")
        if not steps:
            steps.append("Refer to the Register Description chapter for the bring-up sequence specific to this IP.")
        for i, s in enumerate(steps, start=1):
            p = doc.add_paragraph(style="List Number")
            p.add_run(s)

    # 5.2+ Scenarios
    for scn in scenarios:
        name = str(scn.get("name") or scn.get("id") or "Scenario")
        doc.add_heading(name, level=2)
        summary = str(scn.get("summary") or scn.get("description") or "").strip()
        if summary:
            doc.add_paragraph(summary)
        steps = scn.get("steps") if isinstance(scn.get("steps"), list) else []
        rows = []
        for st in steps:
            if not isinstance(st, dict):
                continue
            cycle = st.get("cycle")
            rows.append([
                _ssot_md_scalar(cycle if cycle is not None else ""),
                _ssot_md_scalar(st.get("action") or st.get("description") or ""),
                _ssot_md_scalar(st.get("fl_state") or st.get("function_state") or ""),
                _ssot_md_scalar(st.get("cl_state") or st.get("cycle_state") or st.get("stage") or ""),
            ])
        if rows:
            _ssot_docx_table_from_rows(doc, ["Cycle", "Action", "FL state", "CL state"], rows)

_SSOT_DOCX_APPENDIX_KEYS = (
    "decomposition", "dataflow", "cycle_model", "clock_reset_domains",
    "cdc_requirements", "rdc_requirements", "memory", "interrupts", "fsm",
    "rtl_contract", "timing", "power", "security", "error_handling",
    "debug_observability", "integration", "dft", "synthesis", "pnr",
    "coding_rules", "reuse_modules", "custom", "dir_structure", "filelist",
    "test_requirements", "quality_gates", "traceability", "workflow_todos",
    "generation_flow",
)

def _ssot_to_docx(data: dict, ip: str, out_path: Path) -> None:
    from docx import Document  # type: ignore

    doc = Document()
    safe_data = data if isinstance(data, dict) else {}
    _ssot_docx_set_footer(doc, ip)
    _ssot_docx_apply_heading_numbering(doc)
    _ssot_docx_cover_page(doc, ip, safe_data)
    _ssot_docx_render_revision_history(doc, safe_data)
    _ssot_docx_add_toc(doc)
    _ssot_docx_add_list_of_tables(doc)
    _ssot_docx_add_list_of_figures(doc)

    if not isinstance(data, dict):
        doc.add_heading("Raw Document", level=1)
        _ssot_docx_yaml_block(doc, data)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return

    # ── Chapter 1. Introduction ────────────────────────────────────
    doc.add_heading("Introduction", level=1)
    top = data.get("top_module") if isinstance(data.get("top_module"), dict) else {}
    intro_text = str(top.get("description") or top.get("purpose") or "").strip()
    if intro_text:
        doc.add_paragraph(intro_text)
    if not _ssot_section_is_empty(data.get("top_module")) or not _ssot_section_is_empty(data.get("sub_modules")):
        doc.add_heading("Block Diagram", level=2)
        _ssot_docx_render_block_diagram(doc, data)
    if not _ssot_section_is_empty(data.get("features")):
        doc.add_heading("Features", level=2)
        _ssot_docx_render_features(doc, data["features"])
    if (not _ssot_section_is_empty(data.get("features"))
            or not _ssot_section_is_empty(data.get("sub_modules"))
            or not _ssot_section_is_empty(data.get("function_model"))):
        doc.add_heading("Function Description", level=2)
        _ssot_docx_render_function_description(doc, data)

    # ── Chapter 2. Signal Description ──────────────────────────────
    if not _ssot_section_is_empty(data.get("io_list")):
        _ssot_docx_page_break(doc)
        doc.add_heading("Signal Description", level=1)
        _ssot_docx_render_io_list(doc, data["io_list"])

    # ── Chapter 3. Programming Model ───────────────────────────────
    if not _ssot_section_is_empty(data.get("registers")):
        _ssot_docx_page_break(doc)
        doc.add_heading("Programming Model", level=1)
        _ssot_docx_render_programming_model(doc, data["registers"])

    # ── Chapter 4. Hardware Configuration Options ──────────────────
    if not _ssot_section_is_empty(data.get("parameters")):
        _ssot_docx_page_break(doc)
        doc.add_heading("Hardware Configuration Options", level=1)
        _ssot_docx_render_parameters(doc, data["parameters"])

    # ── Chapter 5. Programming Sequence ────────────────────────────
    cycle_model = data.get("cycle_model") if isinstance(data.get("cycle_model"), dict) else {}
    fn_model = data.get("function_model") if isinstance(data.get("function_model"), dict) else {}
    has_scenarios = (isinstance(cycle_model.get("scenarios"), list) and cycle_model["scenarios"]) \
        or (isinstance(fn_model.get("scenarios"), list) and fn_model["scenarios"])
    if has_scenarios or not _ssot_section_is_empty(data.get("registers")):
        _ssot_docx_page_break(doc)
        doc.add_heading("Programming Sequence", level=1)
        _ssot_docx_render_programming_sequence(doc, data)

    # ── Appendix — Detailed sections ───────────────────────────────
    appendix_items = [
        (k, data.get(k)) for k in _SSOT_DOCX_APPENDIX_KEYS
        if k in data and not _ssot_section_is_empty(data.get(k))
    ]
    if appendix_items:
        _ssot_docx_page_break(doc)
        doc.add_heading("Appendix — Detailed Sections", level=1)
        for key, value in appendix_items:
            label = next((lbl for k, lbl in _SSOT_EXPORT_SECTION_ORDER if k == key), key.replace("_", " ").title())
            doc.add_heading(label, level=2)
            try:
                if key == "fsm":
                    _ssot_docx_render_fsm(doc, value)
                elif key == "interrupts":
                    _ssot_docx_render_interrupts(doc, value)
                elif key == "cycle_model":
                    _ssot_docx_render_cycle_model(doc, value)
                elif key == "error_handling":
                    _ssot_docx_render_error_handling(doc, value)
                else:
                    _ssot_docx_render_section(doc, key, value)
            except Exception as exc:
                err = doc.add_paragraph()
                err.add_run(f"(render error: {exc})").italic = True
                _ssot_docx_yaml_block(doc, value)

    # Catch-all: any keys not in the canonical list nor the appendix.
    canonical = {key for key, _ in _SSOT_EXPORT_SECTION_ORDER}
    chapter1_keys = {"top_module", "features", "sub_modules", "function_model"}
    chapter_keys = chapter1_keys | {"io_list", "registers", "parameters"} | set(_SSOT_DOCX_APPENDIX_KEYS)
    other_keys = [k for k in data.keys()
                  if k not in chapter_keys and k not in canonical
                  and not _ssot_section_is_empty(data.get(k))]
    if other_keys:
        _ssot_docx_page_break(doc)
        doc.add_heading("Other Sections", level=1)
        for key in other_keys:
            doc.add_heading(str(key), level=2)
            _ssot_docx_render_section(doc, key, data.get(key))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ── App factory ────────────────────────────────────────────────────
